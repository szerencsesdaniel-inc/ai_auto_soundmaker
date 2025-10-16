"""
DOCX forgatókönyv feldolgozó modul
Feladata: Microsoft Word .docx fájlokból szöveget kinyerni és feldolgozni.
"""

from docx import Document
import re
from typing import Dict, List, Optional
from script_parser import ScriptParser


class DocxParser(ScriptParser):
    """
    DOCX forgatókönyv parser osztály.
    A ScriptParser osztályból öröklődik, de .docx fájlokat olvas be.
    """
    
    def __init__(self, docx_path: str):
        """
        Inicializálja a DOCX parsert.
        
        Args:
            docx_path: A .docx forgatókönyv fájl elérési útja
        """
        self.script_path = docx_path
        self.characters = {}
        self.metadata = {}
        self.scenes = []
        self._text_content = None
    
    def _extract_text_from_docx(self) -> str:
        """
        Kinyeri a teljes szöveget a .docx fájlból.
        
        Returns:
            str: A dokumentum teljes szövege
        """
        try:
            doc = Document(self.script_path)
            
            # Bekezdések összeállítása
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Csak nem-üres bekezdések
                    paragraphs.append(text)
            
            # Táblázatok kezelése (ha vannak párbeszédek táblázatban)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in paragraphs:
                            paragraphs.append(text)
            
            return '\n'.join(paragraphs)
            
        except Exception as e:
            raise Exception(f"Hiba a .docx fájl olvasásakor: {e}")
    
    def parse(self) -> Dict:
        """
        Feldolgozza a teljes .docx forgatókönyvet.
        
        Returns:
            Dict: Strukturált adatok (metadata, characters, scenes)
        """
        # Szöveg kinyerése a .docx-ből
        self._text_content = self._extract_text_from_docx()
        
        if not self._text_content:
            raise Exception("A .docx fájl üres vagy nem tartalmaz szöveget!")
        
        # A szöveg feldolgozása ugyanúgy, mint a .txt-nél
        self._extract_metadata(self._text_content)
        self._extract_characters(self._text_content)
        self._extract_scenes(self._text_content)
        
        return {
            'metadata': self.metadata,
            'characters': self.characters,
            'scenes': self.scenes
        }
    
    def get_text_content(self) -> str:
        """Visszaadja a kinyert szöveget (debug célra)."""
        return self._text_content or ""


def convert_docx_to_txt(docx_path: str, txt_path: str) -> bool:
    """
    Konvertál egy .docx fájlt .txt-vé (opcionális segédfüggvény).
    
    Args:
        docx_path: Input .docx fájl
        txt_path: Output .txt fájl
        
    Returns:
        bool: Sikeres volt-e a konverzió
    """
    try:
        doc = Document(docx_path)
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
        
        return True
    except Exception as e:
        print(f"Konverziós hiba: {e}")
        return False
